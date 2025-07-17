import streamlit as st
from SE import *
from TE import *
from BE import *
from config import *
import pandas as pd
import matplotlib.pyplot as plt
import docx2txt
from docx import Document
from docx.shared import Inches
import os
import io
import tempfile
from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.utils import get_column_letter
from io import BytesIO

st.title("🤖 RESULT ANALYSIS PRODUCT 💻")

analysis_result, analysis_root_cause, naac_tab = st.tabs(["Result Analysis", "Root Cause Analysis", "NAAC Application"])

if(analysis_result):
    st.title("Result Analysis")
    st.header("1. Upload Files")

    file1 = st.file_uploader("Curent year Excel File", type=["xlsx", "xls"], key = "1")
    file2 = st.file_uploader("Previous year Excel File", type=["xlsx", "xls"], key = "2")

    year = st.selectbox("select year", ["SE", "TE", "BE"])
    semester = st.selectbox("select semester", ["I", "II"])

    if file1 and file2:
        if st.button("Process File"):
            if(year == 'SE'):
                if(semester == "I"):
                    sub = getConfig("SEM-III")
                else:
                    sub = getConfig("SEM-IV")

                with st.spinner("Processing..."):
                    result = SE_analysis(file1, file2, sub)

                st.success("File processed successfully!")
                st.download_button(
                    label="Download Processed File",
                    data=result,
                    file_name="Result_Analysis.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            elif(year == "TE"):
                if(semester == "I"):
                    sub = getConfig("SEM-V")
                else:
                    sub = getConfig("SEM-VI")

                with st.spinner("Processing..."):
                    result = TE_analysis(file1, file2, sub)

                st.success("File processed successfully!")
                st.download_button(
                    label="Download Processed File",
                    data=result,
                    file_name="Result_Analysis.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            else:
                if(semester == "I"):
                    sub = getConfig("SEM-VII")
                else:
                    sub = getConfig("SEM-VIII")

                with st.spinner("Processing..."):
                    result = BE_analysis(file1, file2, sub)

                st.success("Result Analysis File processed successfully!")
                st.download_button(
                    label="Download Processed File",
                    data=result,
                    file_name="Result_Analysis.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

if(analysis_root_cause):
    # Streamlit page setup
    st.set_page_config(page_title="Root Cause Analysis", layout="wide")
    st.title("📘 Failed Students Root Cause Analysis Report Generator")

    # File upload section

    year = st.selectbox("select year", ["TE", "BE"])

    year_file = st.file_uploader("Upload ", year, " Analysis File (.xlsx)", type=['xlsx'])
    fe_file = st.file_uploader("Upload FE Result File (.xls)", type=['xls'])
    dse_file = st.file_uploader("Upload DSE Admission File (.docx)", type=['docx'])

    if year_file and fe_file and dse_file:
        st.success("All files uploaded successfully. Click the button below to run the analysis.")

        if st.button("🚀 Run Root Cause Analysis"):
            # Temporary directory to store charts
            with tempfile.TemporaryDirectory() as temp_dir:
                try:
                    # Load year Data
                    df_year = pd.read_excel(year_file, sheet_name="Failed students")
                    df_year.rename(columns={"Seat No": "Roll No", "name": "Name"}, inplace=True)
                    df_year = df_year.drop_duplicates(subset=["Name"])
                    df_year["Name"] = df_year["Name"].str.lower().str.strip()

                    # Load FE Data
                    df_fe = pd.read_excel(fe_file)
                    df_fe.rename(columns={"name": "Name", "gender": "Gender", "region": "Place of Living",
                                        "seat_type2": "Admission Type", "category": "Category"}, inplace=True)
                    df_fe = df_fe.loc[:, ~df_fe.columns.duplicated()]
                    df_fe["Name"] = df_fe["Name"].str.lower().str.strip()

                    # Load DSE Data
                    dse_text = docx2txt.process(dse_file)
                    dse_lines = dse_text.split("\n")
                    dse_records = []
                    for line in dse_lines:
                        parts = line.split()
                        if len(parts) > 3:
                            name = " ".join(parts[1:-2])
                            category = parts[-2]
                            seat_type = parts[-1]
                            dse_records.append([name.lower().strip(), category, seat_type])
                    df_dse = pd.DataFrame(dse_records, columns=["Name", "Category", "Admission Type"])

                    # Merging
                    merged_fe = df_year.merge(df_fe[["Name", "Gender", "Place of Living", "Admission Type", "Category"]],
                                            on="Name", how="left")
                    merged_dse = df_year.merge(df_dse[["Name", "Admission Type", "Category"]],
                                            on="Name", how="left")
                    merged_dse = merged_dse.merge(df_fe[["Name", "Gender", "Place of Living"]],
                                                on="Name", how="left")

                    combined_df = pd.concat([merged_fe, merged_dse], ignore_index=True).drop_duplicates(subset=["Name"])
                    combined_df["Gender"].fillna("DSE", inplace=True)
                    combined_df["Admission Type"].fillna("DSE", inplace=True)
                    combined_df["Place of Living"].fillna("DSE", inplace=True)
                    combined_df["Category"].fillna("Unknown", inplace=True)

                    # Word Document
                    doc = Document()
                    doc.add_heading("Failed Students Analysis", level=1)

                    # Chart plotting function
                    def plot_bar_and_pie(data, column, title, filename_prefix):
                        if column not in data.columns or data[column].isnull().all():
                            return
                        value_counts = data[column].value_counts()

                        # Bar
                        bar_path = os.path.join(temp_dir, f"{filename_prefix}_bar.png")
                        plt.figure(figsize=(8, 5))
                        bars = plt.bar(value_counts.index, value_counts.values, color="skyblue")
                        plt.title(f"{title} (Bar Chart)")
                        plt.xlabel(column)
                        plt.ylabel("Number of Students")
                        plt.xticks(rotation=45)
                        for bar in bars:
                            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height(),
                                    str(int(bar.get_height())), ha='center', va='bottom', fontsize=10)
                        plt.tight_layout()
                        plt.savefig(bar_path)
                        plt.close()

                        # Pie
                        pie_path = os.path.join(temp_dir, f"{filename_prefix}_pie.png")
                        plt.figure(figsize=(7, 7))
                        plt.pie(value_counts, labels=value_counts.index, autopct='%1.1f%%')
                        plt.title(f"{title} (Pie Chart)")
                        plt.tight_layout()
                        plt.savefig(pie_path)
                        plt.close()

                        doc.add_heading(title, level=2)
                        doc.add_paragraph("Bar Chart:")
                        doc.add_picture(bar_path, width=Inches(5.5))
                        doc.add_paragraph("Pie Chart:")
                        doc.add_picture(pie_path, width=Inches(5.5))

                    # Generate Charts
                    plot_bar_and_pie(combined_df, "Gender", "Failed Students by Gender", "gender")
                    plot_bar_and_pie(combined_df, "Admission Type", "Failed Students by Admission Type", "admission_type")
                    plot_bar_and_pie(combined_df, "Place of Living", "Failed Students by Residence", "place_of_living")
                    plot_bar_and_pie(combined_df, "Category", "Failed Students by Category", "category")

                    # Save final report to buffer
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)

                    st.success("📄 Report generated successfully!")
                    st.download_button("⬇️ Download Word Report",
                                    data=doc_buffer,
                                    file_name="Failed_Students_Analysis.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                except Exception as e:
                    st.error(f"❌ An error occurred: {e}")
    else:
        st.warning("Please upload all three files to start the analysis.")

if(naac_tab):
    # Common formatting functions
    def add_heading(ws, text, row):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        cell = ws.cell(row=row, column=1, value=text)
        cell.font = Font(size=14, bold=True)
        cell.alignment = Alignment(horizontal="center")
        return row + 2

    def write_df(ws, df, start_row):
        for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), start=start_row):
            for c_idx, val in enumerate(row, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=val)
                if r_idx == start_row:
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                else:
                    cell.alignment = Alignment(horizontal="left", vertical="center")
        # Auto-width columns
        for col in ws.iter_cols(min_row=start_row, max_row=ws.max_row):
            max_length = max(len(str(cell.value)) for cell in col) + 2
            ws.column_dimensions[get_column_letter(col[0].column)].width = max_length
        return r_idx + 2

    def generate_report(analysis_path, marks_path, params):
        """Unified report generator with parameterized configuration"""
        # Load and process data
        def load_data(file):
            df = pd.read_excel(file, header=[0, 1])
            df.columns = [f"{c[0]} ({c[1]})" if "Unnamed" not in c[0] else c[1] for c in df.columns]
            return df

        df_analysis = load_data(analysis_path)
        df_marks = load_data(marks_path)

        # 1) Subject-wise percentage table
        percent_dict = df_analysis.iloc[-1].to_dict()
        subject_data = []
        for i, (subj, col, col_type) in enumerate(params['subjects'], 1):
            row = {"S.N.": i, f"SUBJECT({params['semester']})": subj, 
                **{f"{t} %": "-" for t in ['TH', 'TW', 'PR', 'OR']}}
            if col in percent_dict and pd.notna(percent_dict[col]):
                row[f"{col_type} %"] = f"{percent_dict[col]:.2f}"
            subject_data.append(row)
        subject_df = pd.DataFrame(subject_data)


        # 2) Overall result summary
        total_students = int(df_analysis.iloc[0, -1])
        sgpa_data = {
            "ALL CLEAR": df_analysis.iloc[18, -1],
            "DISTINCTION (> 7.75 SGPA)": df_analysis.iloc[5, -1],
            "FIRST CLASS (6.75 TO 7.74 SGPA)": df_analysis.iloc[7, -1],
            "HIGH.SECOND CLASS (6.25 TO 6.74 SGPA)": df_analysis.iloc[9, -1],
            "SECOND CLASS (5.5 TO 6.24 SGPA)": df_analysis.iloc[11, -1],
            "PASS CLASS (4.0 TO 5.49 SGPA)": df_analysis.iloc[13, -1],
            "FAIL": df_analysis.iloc[15, -1]
        }
        overall_df = pd.DataFrame([(k, round(v, 2)) for k, v in sgpa_data.items()], 
                                columns=["RESULT", "NO OF STUDENTS (%)"])

        # 3) Class toppers
        df_marks.rename(columns={params['name_col']: "Name", params['sgpa_col']: "SGPA"}, inplace=True)
        df_marks["SGPA"] = pd.to_numeric(df_marks["SGPA"], errors='coerce').round(2)
        df_sorted = df_marks[["Name", "SGPA"]].dropna().sort_values("SGPA", ascending=False)
        
        top_5_sgpas = df_sorted["SGPA"].drop_duplicates().head(5)
        topper_rows = []
        for rank, sgpa in enumerate(top_5_sgpas, 1):
            for idx, name in enumerate(df_sorted[df_sorted["SGPA"] == sgpa]["Name"]):
                topper_rows.append([rank if idx == 0 else "", name, sgpa])
        topper_df = pd.DataFrame(topper_rows, columns=["Rank", "Name of Student", "SGPA"])

        # 4) Subject toppers
        subject_topper_data = []
        for sn, (col, subj) in enumerate(params['subject_map'].items(), 1):
            df_marks[col] = pd.to_numeric(df_marks[col], errors='coerce')
            if (max_marks := df_marks[col].max()) and not pd.isna(max_marks):
                students = df_marks[df_marks[col] == max_marks]["Name"].dropna().tolist()
                for idx, name in enumerate(students):
                    subject_topper_data.append([
                        sn if idx == 0 else "", 
                        subj if idx == 0 else "", 
                        name, 
                        int(max_marks)
                    ])
        subject_topper_df = pd.DataFrame(subject_topper_data, 
            columns=["S.N.", f"SUBJECT ({params['semester']})", "NAME OF THE STUDENT(S)", "MARKS OBTAINED (OUT OF 100)"])

        # Create Excel report
        wb = Workbook()
        ws = wb.active
        ws.title = params['report_title']
        current_row = 1

        for section in [
            ("SUBJECT-WISE % RESULT", subject_df),
            ("OVERALL RESULT", overall_df, f"TOTAL NO OF STUDENTS APPEARED = {total_students}"),
            ("CLASS TOPPERS", topper_df),
            ("SUBJECT TOPPERS", subject_topper_df)
        ]:
            current_row = add_heading(ws, section[0], current_row)
            if section[0] == "OVERALL RESULT":
                ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=6)
                ws.cell(row=current_row, column=1, value=section[2]).alignment = Alignment(horizontal="center")
                current_row += 2
            current_row = write_df(ws, section[1], current_row)

        return wb

    # Streamlit UI

    st.title("NAAC Report Generator 📊")
    st.header("1. Upload Files")

    analysis_file = st.file_uploader("Analysis Excel", type=["xlsx"])
    marks_file = st.file_uploader("Result Excel", type=["xlsx"])

    semester = st.selectbox("2. Select Semester", ["SEM-III", "SEM-IV", "SEM-V", "SEM-VI", "SEM-VII", "SEM-VIII"])

    # In the Streamlit UI section where you create the download button:
    # Enhanced filename generation in download button
    if st.button("Generate Report") and analysis_file and marks_file:
        with st.spinner("Generating..."):
            try:
                config = CONFIG[semester]
                wb = generate_report(analysis_file, marks_file, config)

                # Generate filename with proper prefix and date
                prefix = config['report_title'].split('_')[0]  # SE/TE/BE
                date_code = 'NOV23' if 'NOV' in config['report_title'] else 'May24'

                output = BytesIO()
                wb.save(output)

                st.success("Done!").download_button(
                    "Download Report",
                    output.getvalue(),
                    f"{prefix}_NAAC_{date_code}.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            except Exception as e:
                st.error(f"Error generating report: {str(e)}")

    st.markdown("### Instructions\n1. Upload both files\n2. Select semester\n3. Generate and download")